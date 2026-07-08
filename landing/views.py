from django.shortcuts import render, redirect
from news.models import News
from landing.models import CourseDirection, Review, Expert, MinstroyProgram, Qualification, UnifiedRequest, ThreedGalleryImage, NocPreparationDirection, NocMailSettings, GratitudeItem
from landing.forms import CallbackForm, NOCRequestForm
from django.http import JsonResponse, HttpResponse
from django.core.paginator import Paginator
from django.conf import settings
from django.http import Http404
from django.views.decorators.http import require_POST
from django.urls import reverse
from datetime import datetime
from django.core.mail import EmailMessage
from zipfile import ZipFile, ZIP_DEFLATED

def get_noc_email_recipients():
    mail_settings = NocMailSettings.objects.first()

    if mail_settings:
        to = mail_settings.parse_emails(mail_settings.to_emails)
        cc = mail_settings.parse_emails(mail_settings.cc_emails)
    else:
        to = getattr(settings, "NOC_UPLOAD_TO", []) or []
        cc = getattr(settings, "NOC_UPLOAD_CC", []) or []

    return to, cc

def index(request):
    if request.method == "POST":
        form = CallbackForm(request.POST)
        if form.is_valid():
            form.save()
        return redirect(request.path)

    reviews = Review.objects.all()[:6]
    popular_courses = CourseDirection.objects.filter(featured=True)[:3] or CourseDirection.objects.all()[:3]
    latest_news = News.objects.all().order_by('-created_at')[:3]

    context = {
        'latest_news': latest_news,
        'popular_courses': popular_courses,
        'reviews': reviews
    }

    return render(request, 'landing/index.html', context)

def contact(request):
    return render(request, 'landing/contact.html')



def callback_request(request):
    if request.method == "POST":
        form = CallbackForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.request_type = "callback"
            obj.save()

            if request.headers.get("x-requested-with") == "XMLHttpRequest":
                return JsonResponse({"success": True})

            return HttpResponse("OK")

        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": False, "errors": form.errors}, status=400)

    return JsonResponse({"error": "Invalid method"}, status=405)



def nok_page(request):
    experts = Expert.objects.all()
    qualifications = Qualification.objects.order_by("created_at")[:3]

    if request.method == "POST":
        form = NOCRequestForm(request.POST, request.FILES)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.request_type = "noc"
            obj.save()

            if request.headers.get("x-requested-with") == "XMLHttpRequest":
                return JsonResponse({"success": True})

            return redirect("nok")

        if request.headers.get("x-requested-with") == "XMLHttpRequest":
            return JsonResponse({"success": False, "errors": form.errors})

    else:
        form = NOCRequestForm()

    return render(
        request,
        "landing/nok.html",
        {"experts": experts, "form": form, "qualifications": qualifications}
    )


def prep_request(request):
    if request.method == "POST":
        req_type = request.POST.get("request_type")

        if req_type not in ("prep_expert", "prep_specialist"):
            return JsonResponse({"success": False, "error": "Invalid type"}, status=400)

        name = request.POST.get("name", "").strip()
        phone = request.POST.get("phone", "").strip()
        email = request.POST.get("email", "").strip()
        message = request.POST.get("message", "").strip()

        if not (name and phone and email):
            return JsonResponse({"success": False, "error": "Missing fields"}, status=400)

        UnifiedRequest.objects.create(
            name=name,
            phone=phone,
            email=email,
            message=message,
            request_type=req_type
        )

        return JsonResponse({"success": True})

    return JsonResponse({"error": "Invalid method"}, status=405)


# def minstroy_page(request):
#     programs = MinstroyProgram.objects.all().order_by("id")[:4]
#     form = NOCRequestForm()
#     context = {
#         "programs": programs,
#         "form": form,
#     }
#     return render(request, "landing/minstroy.html", context)

# def minstroy_list(request):
#     programs = MinstroyProgram.objects.all().order_by("id")
#     paginator = Paginator(programs, 14)
#     page_number = request.GET.get("page")
#     page_obj = paginator.get_page(page_number)
#     return render(request, "landing/minstroy_list.html", {"page_obj": page_obj})

def exam(request):
    qs = MinstroyProgram.objects.all()
    fields = {f.name for f in MinstroyProgram._meta.get_fields()}

    if "is_active" in fields:
        qs = qs.filter(is_active=True)
    if "order" in fields:
        qs = qs.order_by("order", "id")
    else:
        qs = qs.order_by("id")

    programs = qs[:4]
    qualifications = Qualification.objects.order_by("created_at")[:3]
    directions = NocPreparationDirection.objects.filter(is_active=True).order_by("kind", "id")
    return render(request, "landing/exam.html", {
        "programs": programs,
        "qualifications": qualifications,
        "directions": directions,
    })



def qualifications_list(request):
    q = (request.GET.get("q") or "").strip()
    page_number = request.GET.get("page")

    base_qs = Qualification.objects.all().order_by("code")

    if q:
        q_cf = q.casefold()

        items = [
            obj for obj in base_qs
            if q_cf in (obj.code or "").casefold()
            or q_cf in (obj.title or "").casefold()
        ]

        paginator = Paginator(items, 12)
        page_obj = paginator.get_page(page_number)
        total = len(items)
    else:
        paginator = Paginator(base_qs, 12)
        page_obj = paginator.get_page(page_number)
        total = base_qs.count()

    return render(request, "landing/qualifications_list.html", {
        "page_obj": page_obj,
        "q": q,
        "total": total,
    })

from io import BytesIO
from pathlib import Path
from docxtpl import DocxTemplate
from collections import defaultdict
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CHECKED = "☑"
UNCHECKED = "☐"


def _p(cell, text, bold=False, center=False, size=10):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def build_auditors_table(subdoc, selected_dirs):
    auditors = [d for d in selected_dirs if d.kind == "auditor"]
    if not auditors:
        subdoc.add_paragraph("Аудиторы: не выбрано.")
        return subdoc

    # Таблица: № | Наименование | отметка | Код проф. квалиф.
    t = subdoc.add_table(rows=1, cols=4)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    _p(hdr[0], "№", bold=True, center=True)
    _p(hdr[1], "Наименование ОПО", bold=True, center=True)
    _p(hdr[2], "", bold=True, center=True)
    _p(hdr[3], "Код проф. квалиф.", bold=True, center=True)

    for i, d in enumerate(auditors, start=1):
        row = t.add_row().cells
        _p(row[0], str(i), center=True)
        _p(row[1], d.title)
        _p(row[2], CHECKED, center=True, size=14)  # выбран — значит отмечен
        _p(row[3], d.code, center=True)

    return subdoc


def build_experts_table(subdoc, selected_dirs):
    experts = [d for d in selected_dirs if d.kind == "expert"]
    if not experts:
        subdoc.add_paragraph("Эксперты: не выбрано.")
        return subdoc

    # группируем по "Наименование ОПО" (title)
    # внутри — по track (TU/ZS) и category (1/2/3)
    grouped = defaultdict(list)
    for d in experts:
        grouped[d.title].append(d)

    # какие категории вообще выбраны (чтобы колонки были только по ним)
    cats = sorted({d.category for d in experts if d.category is not None})
    # если вдруг category не заполнена, можно fallback:
    if not cats:
        cats = [1, 2, 3]  # или [] — но таблица станет пустой

    # Таблица: № | Наименование ОПО | ТУ/ЗС | (Категории...)
    cols = 3 + len(cats)
    t = subdoc.add_table(rows=1, cols=cols)
    t.style = "Table Grid"

    # Header
    hdr = t.rows[0].cells
    _p(hdr[0], "№", bold=True, center=True)
    _p(hdr[1], "Наименование ОПО", bold=True, center=True)
    _p(hdr[2], "ТУ/ЗС", bold=True, center=True)
    for j, cat in enumerate(cats):
        _p(hdr[3 + j], f"Категория, Код", bold=True, center=True)

    # Rows: по каждому ОПО делаем 2 строки (ТУ и ЗС) только если они есть среди выбранных
    # (если выбран только ТУ — будет 1 строка)
    idx = 1
    for title, items in grouped.items():
        # map: track -> {cat -> dir}
        by_track = defaultdict(dict)
        for d in items:
            by_track[d.track][d.category] = d

        tracks_order = []
        if "TU" in by_track: tracks_order.append("TU")
        if "ZS" in by_track: tracks_order.append("ZS")

        first_row_for_title = True
        for track in tracks_order:
            row = t.add_row().cells

            if first_row_for_title:
                _p(row[0], str(idx), center=True)
                _p(row[1], title)
                first_row_for_title = False
                idx += 1
            else:
                _p(row[0], "")
                _p(row[1], "")

            _p(row[2], "ТУ" if track == "TU" else "ЗС", center=True)

            for j, cat in enumerate(cats):
                d = by_track.get(track, {}).get(cat)
                if d:
                    # в клетке: код + чекбокс (чекбокс отмечен, потому что это выбранное направление)
                    _p(row[3 + j], f"Категория {cat}\nКод {d.code}\n{CHECKED}", center=True)
                else:
                    # категории нет среди выбранных для этой строки — можно пусто
                    _p(row[3 + j], "", center=True)

    return subdoc


@require_POST
def noc_exam_print(request):
    def g(name):
        return (request.POST.get(name) or "").strip()

    applicant_type = g("applicant_type")
    candidate_fio = g("candidate_fio")
    candidate_phone = g("candidate_phone")
    candidate_email = g("candidate_email")

    errors = {}
    if applicant_type not in ("company", "person"):
        errors["applicant_type"] = "Выберите тип заявителя"
    if not candidate_fio:
        errors["candidate_fio"] = "Укажите ФИО"
    if not candidate_phone:
        errors["candidate_phone"] = "Укажите телефон"
    if not candidate_email:
        errors["candidate_email"] = "Укажите email"

    birth_date_raw = g("birth_date")
    birth_date = ""
    if birth_date_raw:
        try:
            birth_date = datetime.strptime(birth_date_raw, "%Y-%m-%d").date().strftime("%d.%m.%Y")
        except ValueError:
            errors["birth_date"] = "Неверный формат даты"

    if errors:
        return JsonResponse({"success": False, "errors": errors}, status=400)

    # ---- направления ----

    direction_ids = request.POST.getlist("directions")
    selected_dirs = list(NocPreparationDirection.objects.filter(id__in=direction_ids)) if direction_ids else []

    edo_enabled = (g("edo_enabled") == "1")
    today = datetime.now().strftime("%d.%m.%Y")

# заявка
    request_tpl_path = Path(settings.BASE_DIR) / "landing" / "docx_templates" / "noc_request_template.docx"
    request_doc = DocxTemplate(str(request_tpl_path))

    experts_subdoc = request_doc.new_subdoc()
    build_experts_table(experts_subdoc, selected_dirs)

    auditors_subdoc = request_doc.new_subdoc()
    build_auditors_table(auditors_subdoc, selected_dirs)

    request_context = {
        "today": today,
        "applicant_type_label": "Предприятие-плательщик" if applicant_type == "company" else "Частное лицо",

        "applicant_name": g("applicant_name"),
        "contacts": g("contacts"),
        "legal_address": g("legal_address"),
        "postal_address": g("postal_address"),

        "edo_enabled_label": "Да" if edo_enabled else "Нет",
        "edo_service": g("edo_service") if edo_enabled else "",
        "edo_yes": "☑" if edo_enabled else "☐",
        "edo_no": "☐" if edo_enabled else "☑",

        "rs": g("rs"),
        "bank": g("bank"),
        "ks": g("ks"),
        "inn": g("inn"),
        "kpp": g("kpp"),
        "bik": g("bik"),
        "okpo": g("okpo"),
        "okved": g("okved"),

        "candidate_fio": candidate_fio,
        "birth_date": birth_date,
        "position": g("position"),
        "residence": g("residence"),
        "passport_data": g("passport_data"),
        "candidate_phone": candidate_phone,
        "candidate_email": candidate_email,

        "contact_fio": g("contact_fio"),
        "contact_phone": g("contact_phone"),
        "contact_email": g("contact_email"),
        "comment": g("comment"),

        "experts_table": experts_subdoc,
        "auditors_table": auditors_subdoc,
    }

    request_doc.render(request_context)

    request_buff = BytesIO()
    request_doc.save(request_buff)
    request_bytes = request_buff.getvalue()

# согласие
    consent_tpl_path = Path(settings.BASE_DIR) / "landing" / "docx_templates" / "noc_consent_template.docx"
    consent_doc = DocxTemplate(str(consent_tpl_path))

    consent_context = {
        "candidate_fio": candidate_fio,
        "passport_data": g("passport_data"),
        "residence": g("residence"),
        "today": today,
    }

    consent_doc.render(consent_context)

    consent_buff = BytesIO()
    consent_doc.save(consent_buff)
    consent_bytes = consent_buff.getvalue()

# упаковка в zip
    zip_buff = BytesIO()

    request_filename = f"zayavka-nok-{datetime.now():%Y-%m-%d}.docx"
    consent_filename = f"soglasie-na-obrabotku-pd-{datetime.now():%Y-%m-%d}.docx"
    zip_filename = f"dokumenty-nok-{datetime.now():%Y-%m-%d}.zip"

    with ZipFile(zip_buff, "w", compression=ZIP_DEFLATED) as zf:
        zf.writestr(request_filename, request_bytes)
        zf.writestr(consent_filename, consent_bytes)

    zip_buff.seek(0)

    resp = HttpResponse(
        zip_buff.getvalue(),
        content_type="application/zip",
    )
    resp["Content-Disposition"] = f"attachment; filename*=UTF-8''{zip_filename}"
    return resp

# # приём подписанной заявки НОК
# @require_POST
# def noc_exam_upload(request):
#     file = request.FILES.get("file")
#     comment = (request.POST.get("comment") or "").strip()

#     errors = {}
#     if not file:
#         errors["file"] = "Прикрепите файл подписанной заявки"

#     if errors:
#         return JsonResponse({"success": False, "errors": errors}, status=400)

#     UnifiedRequest.objects.create(
#       request_type="noc_signed",
#       name="Подписанная заявка (НОК)",
#       message=comment,
#       file=file,
#     )

#     return JsonResponse({"success": True})


def attestation(request):
    return render(request, 'landing/attestation.html')

def TD(request):
    form = NOCRequestForm()
    context = {
        "form": form,
        "page": "TD",
    }
    return render(request, 'landing/TD.html', context)

def IO(request):
    form = NOCRequestForm()
    context = {
        "form": form,
        "page": "IO",
    }
    return render(request, 'landing/TD.html', context)

def NK(request):
    form = NOCRequestForm()
    context = {
        "form": form,
        "page": "NK",
    }
    return render(request, 'landing/TD.html', context)

def docs(request):
    return render(request, 'landing/docs.html')

def preparation(request):
    return render(request, 'landing/preparation.html')

def safety(request):
    return render(request, 'landing/safety.html')

def threed(request):
    if not settings.LANDING_FEATURES.get("ENABLE_THREED_PAGE", True):
        raise Http404
    gallery = ThreedGalleryImage.objects.all()
    form = NOCRequestForm()
    context = {
        'gallery': gallery,
        "form": form,
    }
    return render(request, 'landing/threed.html', context)


ABOUT_EDUCATION_SLUGS = {
    "basic-information",
    "structure",
    "documents",
    "education",
    "leadership",
    "paid-services",
    "international-cooperation",
    "material-and-technical-support",
    "vacancies",
    "financial-activity",
}

def about_education_index(request):
    return redirect("about_education_page", slug="basic-information")


def about_education_page(request, slug):
    if slug not in ABOUT_EDUCATION_SLUGS:
        raise Http404

    return render(
        request,
        "landing/about_education.html",
        {
            "active_slug": slug,
        },
    )

def about_complaints(request):
    return render(request, "landing/about_complaints.html")


def about_impartiality(request):
    return render(request, "landing/about_impartiality.html")


def about_gratitude(request):
    gratitude_items = GratitudeItem.objects.filter(is_active=True).order_by("order", "-date", "-created_at")
    return render(
        request,
        "landing/about_gratitude.html",
        {"gratitude_items": gratitude_items},
    )


MAX_UPLOAD_MB = 20

@require_POST
def noc_exam_upload(request):
    uploaded = request.FILES.get("file")
    consent_uploaded = request.FILES.get("consent_file")
    comment = (request.POST.get("comment") or "").strip()

    errors = {}
    if not uploaded:
        errors["file"] = "Прикрепите файл подписанной заявки"

    if not consent_uploaded:
        errors["consent_file"] = "Прикрепите файл согласия на обработку персональных данных"

    if uploaded and uploaded.size > MAX_UPLOAD_MB * 1024 * 1024:
        errors["file"] = f"Файл слишком большой. Максимум {MAX_UPLOAD_MB} МБ."

    if consent_uploaded and consent_uploaded.size > MAX_UPLOAD_MB * 1024 * 1024:
        errors["consent_file"] = f"Файл слишком большой. Максимум {MAX_UPLOAD_MB} МБ."

    if errors:
        return JsonResponse({"success": False, "errors": errors}, status=400)

    try:
        saved_request = UnifiedRequest.objects.create(
            request_type="noc_signed",
            name="Загрузка подписанной заявки НОК",
            message=comment,
            file=uploaded,
            consent_file=consent_uploaded,
        )
    except Exception:
        return JsonResponse(
            {"success": False, "errors": {"file": "Не удалось сохранить заявку"}},
            status=500
        )

    # защита от ситуации, когда забыли настроить SMTP
    if not getattr(settings, "EMAIL_HOST_USER", "") or not getattr(settings, "EMAIL_HOST_PASSWORD", ""):
        return JsonResponse(
            {
                "success": True,
                "warning": "Заявка сохранена, но почта не настроена."
            }
        )

    subject = "Подписанная заявка на подготовку к НОК"
    body = "\n".join([
        "Здравствуйте!",
        "",
        "Поступила подписанная заявка на подготовку к НОК с сайта ucbp.ru.",
        "",
        f"Комментарий пользователя: {comment or '—'}",
        "",
        f"Файл заявки: {uploaded.name}",
        f"Размер файла заявки: {uploaded.size} байт",
        f"Файл согласия: {consent_uploaded.name}",
        f"Размер файла согласия: {consent_uploaded.size} байт",
        f"ID заявки в базе: {saved_request.id}",
        "",
        "Примечание: некоторые почтовые сервисы не отображают таблицы в предпросмотре Word-документов корректно. Пожалуйста, скачайте файл — в скачанном документе таблица направлений отображается правильно.",
    ])

    from_email = getattr(settings, "DEFAULT_FROM_EMAIL", "") or getattr(settings, "EMAIL_HOST_USER", "")
    to, cc = get_noc_email_recipients()

    if not to:
        return JsonResponse(
            {
                "success": True,
                "warning": "Заявка сохранена, но не указаны получатели письма."
            }
        )

    msg = EmailMessage(
        subject=subject,
        body=body,
        from_email=from_email,
        to=to,
        cc=cc,
    )

    uploaded.seek(0)
    consent_uploaded.seek(0)

    msg.attach(
        uploaded.name,
        uploaded.read(),
        uploaded.content_type or "application/octet-stream"
    )

    msg.attach(
        consent_uploaded.name,
        consent_uploaded.read(),
        consent_uploaded.content_type or "application/octet-stream"
    )

    try:
        msg.send(fail_silently=False)
    except Exception:
        return JsonResponse(
            {
                "success": True,
                "warning": "Заявка сохранена, но письмо отправить не удалось."
            }
        )

    return JsonResponse({"success": True})

